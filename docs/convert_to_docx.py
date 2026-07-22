#!/usr/bin/env python3
"""
Convert Markdown documentation files to DOCX format
Requires: pip install python-docx markdown
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown(md_content):
    """Parse markdown content into structured sections."""
    lines = md_content.split('\n')
    sections = []
    current_section = []
    
    for line in lines:
        if line.startswith('# ') and not line.startswith('## '):
            # Save previous section if exists
            if current_section:
                sections.append(current_section)
            current_section = [line]
        else:
            current_section.append(line)
    
    # Add last section
    if current_section:
        sections.append(current_section)
    
    return sections

def add_heading(doc, text, level):
    """Add a heading to the document."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt({1: 16, 2: 14, 3: 12}.get(level, 12))
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, code=False):
    """Add a paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if code:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    return p

def add_code_block(doc, code_text):
    """Add a code block to the document."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_from_markdown(doc, table_text):
    """Add a table from markdown format."""
    lines = table_text.strip().split('\n')
    if len(lines) < 3:
        return
    
    # Parse header
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    
    # Skip separator line (lines[1])
    # Parse data rows
    rows = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
    
    # Create table
    if headers and rows:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Add header
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            if i < len(hdr_cells):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = cell_text
        
        doc.add_paragraph()  # Add space after table

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX format."""
    print(f"Converting {md_file} to {docx_file}...")
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content into lines
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Main title (H1)
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line.replace('# ', '').strip()
            add_heading(doc, title_text, level=1)
            i += 1
            continue
        
        # H2 headings
        if line.startswith('## '):
            heading_text = line.replace('## ', '').strip()
            add_heading(doc, heading_text, level=2)
            i += 1
            continue
        
        # H3 headings
        if line.startswith('### '):
            heading_text = line.replace('### ', '').strip()
            add_heading(doc, heading_text, level=3)
            i += 1
            continue
        
        # H4 headings
        if line.startswith('#### '):
            heading_text = line.replace('#### ', '').strip()
            add_heading(doc, heading_text, level=4)
            i += 1
            continue
        
        # Code blocks
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, '\n'.join(code_lines))
            i += 1
            continue
        
        # Tables
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            add_table_from_markdown(doc, '\n'.join(table_lines))
            continue
        
        # Bullet points
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            # Handle bold and italic
            bullet_text = re.sub(r'\*\*(.+?)\*\*', r'\1', bullet_text)
            bullet_text = re.sub(r'\*(.+?)\*', r'\1', bullet_text)
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.\s*', '', line)
            list_text = re.sub(r'\*\*(.+?)\*\*', r'\1', list_text)
            p = doc.add_paragraph(list_text, style='List Number')
            i += 1
            continue
        
        # Regular paragraphs
        if line.strip():
            # Clean markdown formatting
            text = line.strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.+?)`', r'\1', text)  # Inline code
            
            # Check if it's a note/warning/tip block
            if text.startswith('**') and '**' in text[2:]:
                parts = text.split('**')
                if len(parts) >= 3:
                    p = doc.add_paragraph()
                    p.add_run(parts[1]).bold = True
                    if len(parts) > 3:
                        p.add_run(''.join(parts[3:]))
                else:
                    add_paragraph(doc, text)
            else:
                add_paragraph(doc, text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ Successfully created {docx_file}")

def main():
    """Main conversion function."""
    # Files to convert
    files = [
        ('docs/Deployment_Documentation.md', 'docs/Deployment_Documentation.docx'),
        ('docs/Configuration_Management_Documentation.md', 'docs/Configuration_Management_Documentation.docx')
    ]
    
    print("=" * 60)
    print("NimbusMart Documentation Converter")
    print("Converting Markdown to DOCX format")
    print("=" * 60)
    print()
    
    for md_file, docx_file in files:
        if os.path.exists(md_file):
            try:
                convert_md_to_docx(md_file, docx_file)
            except Exception as e:
                print(f"✗ Error converting {md_file}: {e}")
        else:
            print(f"✗ File not found: {md_file}")
    
    print()
    print("=" * 60)
    print("Conversion complete!")
    print("=" * 60)
    print()
    print("Generated files:")
    for _, docx_file in files:
        if os.path.exists(docx_file):
            size = os.path.getsize(docx_file)
            print(f"  - {docx_file} ({size:,} bytes)")
    
    print()
    print("Alternative conversion methods if needed:")
    print("1. Online: https://cloudconvert.com/md-to-docx")
    print("2. VS Code: Install 'Markdown All in One' extension")
    print("3. Pandoc: pandoc input.md -o output.docx")

if __name__ == '__main__':
    main()